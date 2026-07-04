# -*- coding: utf-8 -*-
"""
Проверка работ учеников (математика, биология, химия, английский)
Streamlit + Claude API (поддержка реселлерских эндпоинтов через Base URL).

Запуск:  streamlit run app.py
"""

import base64
import io
import json
import re
import time

import httpx
import pandas as pd
import streamlit as st
from PIL import Image

try:
    import docx  # python-docx
except ImportError:
    docx = None

try:
    import fitz  # PyMuPDF — нужен для PDF в OpenAI-совместимом режиме
except ImportError:
    fitz = None

# ============================== НАСТРОЙКИ ==============================

APP_TITLE = "Проверка работ учеников"

SUBJECTS = {
    "Математика": "математика (арифметика, алгебра, геометрия, начала анализа)",
    "Биология": "биология",
    "Химия": "химия",
    "Английский язык": "английский язык (грамматика, лексика, письмо)",
}

DEFAULT_MODELS = [
    "claude-sonnet-4-6",
    "claude-opus-4-8",
    "claude-haiku-4-5-20251001",
]

MAX_IMAGE_SIDE = 1568          # рекомендация API по размеру изображений
MAX_IMAGE_BYTES = 4_500_000    # запас до лимита 5 МБ на изображение
API_VERSION = "2023-06-01"

# ============================ СИСТЕМНЫЙ ПРОМПТ ============================

SYSTEM_PROMPT = """Ты — опытный школьный преподаватель и эксперт по проверке письменных работ.
Твоя задача: внимательно проверить работу ученика по предмету «{subject}» ({subject_hint}),
уровень: {grade_level}. Строгость проверки: {strictness}.

Правила проверки:
1. Разбери работу по заданиям. Если нумерации нет — раздели логически и пронумеруй сам.
2. Для каждого задания определи: что ответил ученик, верно ли это (correct / partial / wrong),
   тип ошибки (если есть), краткий комментарий и правильное решение/ответ.
3. Проверяй ХОД РЕШЕНИЯ, а не только ответ: арифметические ошибки, логические, ошибки в формулах,
   в терминологии, орфографические/грамматические (для английского — grammar, vocabulary, spelling, word order).
4. Если часть работы неразборчива — пометь задание как "unreadable" в error_type и напиши об этом в комментарии, не выдумывай содержание.
5. Пиши комментарии по-русски, доброжелательно, но конкретно. Примеры на английском (для англ. языка) оставляй на английском.
6. Оценку ставь по 100-балльной шкале и дублируй в 5-балльной ({grading_note}).

Ответ верни СТРОГО в виде одного JSON-объекта без пояснений до или после, без markdown-ограждений.
Схема JSON:
{{
  "student_name": "имя ученика, если указано в работе, иначе null",
  "work_title": "название/тема работы, если понятна, иначе null",
  "overall_score": <число 0-100>,
  "grade_5": <число 2-5>,
  "summary": "развернутый общий вывод о работе, 3-6 предложений",
  "tasks": [
    {{
      "number": "номер или условное название задания",
      "topic": "тема/навык",
      "student_answer": "кратко: что написал ученик",
      "verdict": "correct" | "partial" | "wrong" | "skipped",
      "score": <набранные баллы>,
      "max_score": <максимум баллов>,
      "error_type": "тип ошибки или null (напр.: вычислительная, логическая, формула, терминология, grammar, spelling, vocabulary, unreadable)",
      "comment": "комментарий преподавателя к заданию",
      "correct_solution": "правильное решение или ответ (кратко)"
    }}
  ],
  "strengths": ["сильные стороны работы", "..."],
  "weaknesses": ["слабые места и пробелы", "..."],
  "recommendations": ["конкретные рекомендации: что повторить, какие упражнения сделать", "..."],
  "teacher_comment": "итоговый развернутый комментарий для ученика (обращение на «ты», поддерживающий тон)"
}}"""


# ============================ ОБРАБОТКА ФАЙЛОВ ============================

def prepare_image(file_bytes: bytes) -> tuple[str, str]:
    """Приводит изображение к допустимому размеру, возвращает (base64, media_type)."""
    img = Image.open(io.BytesIO(file_bytes))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    w, h = img.size
    scale = min(1.0, MAX_IMAGE_SIDE / max(w, h))
    if scale < 1.0:
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    quality = 90
    while True:
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality)
        data = buf.getvalue()
        if len(data) <= MAX_IMAGE_BYTES or quality <= 40:
            break
        quality -= 10

    return base64.b64encode(data).decode(), "image/jpeg"


def extract_docx_text(file_bytes: bytes) -> str:
    """Извлекает текст из .docx, включая таблицы."""
    if docx is None:
        raise RuntimeError("Библиотека python-docx не установлена. Выполните: pip install python-docx")
    document = docx.Document(io.BytesIO(file_bytes))
    parts = []
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def pdf_to_image_blocks(data: bytes, name: str, max_pages: int = 20) -> list:
    """Рендерит страницы PDF в JPEG-блоки (для эндпоинтов без нативной поддержки PDF)."""
    if fitz is None:
        raise RuntimeError("Для PDF в OpenAI-режиме нужен PyMuPDF: pip install pymupdf")
    blocks = []
    pdf = fitz.open(stream=data, filetype="pdf")
    for i, page in enumerate(pdf):
        if i >= max_pages:
            blocks.append({"type": "text",
                           "text": f"[{name}: показаны первые {max_pages} страниц]"})
            break
        pix = page.get_pixmap(dpi=150)
        b64, media = prepare_image(pix.tobytes("jpeg"))
        blocks.append({"type": "text", "text": f"[{name}, страница {i + 1}]"})
        blocks.append({"type": "image",
                       "source": {"type": "base64", "media_type": media, "data": b64}})
    pdf.close()
    return blocks


def build_content_blocks(uploaded_files, api_format: str = "anthropic") -> tuple[list, list[str]]:
    """Собирает content-блоки для API из загруженных файлов. Возвращает (blocks, warnings)."""
    blocks, warnings = [], []
    for f in uploaded_files:
        name = f.name
        data = f.getvalue()
        ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""

        if ext in ("jpg", "jpeg", "png", "webp", "gif", "bmp", "tiff"):
            try:
                b64, media = prepare_image(data)
                blocks.append({"type": "text", "text": f"[Фото работы: {name}]"})
                blocks.append({
                    "type": "image",
                    "source": {"type": "base64", "media_type": media, "data": b64},
                })
            except Exception as e:
                warnings.append(f"Не удалось обработать изображение {name}: {e}")

        elif ext == "pdf":
            if len(data) > 30_000_000:
                warnings.append(f"PDF {name} больше 30 МБ — пропущен.")
                continue
            if api_format == "openai":
                # OpenAI-совместимые эндпоинты не принимают PDF — рендерим страницы в картинки
                try:
                    blocks.extend(pdf_to_image_blocks(data, name))
                except Exception as e:
                    warnings.append(f"Не удалось обработать PDF {name}: {e}")
            else:
                blocks.append({"type": "text", "text": f"[PDF-файл работы: {name}]"})
                blocks.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": base64.b64encode(data).decode(),
                    },
                })

        elif ext == "docx":
            try:
                text = extract_docx_text(data)
                if not text.strip():
                    warnings.append(f"В {name} не найден текст.")
                else:
                    blocks.append({"type": "text",
                                   "text": f"[Текст из документа {name}]:\n\n{text}"})
            except Exception as e:
                warnings.append(f"Не удалось прочитать {name}: {e}")

        elif ext == "doc":
            warnings.append(f"Формат .doc не поддерживается ({name}). Пересохраните файл как .docx.")
        else:
            warnings.append(f"Файл {name}: неподдерживаемый формат ({ext}).")

    return blocks, warnings


# ============================== ВЫЗОВ API ==============================

def to_openai_content(blocks: list) -> list:
    """Конвертирует внутренние блоки (формат Anthropic) в формат OpenAI chat completions."""
    out = []
    for b in blocks:
        if b["type"] == "text":
            out.append({"type": "text", "text": b["text"]})
        elif b["type"] == "image":
            src = b["source"]
            out.append({
                "type": "image_url",
                "image_url": {"url": f"data:{src['media_type']};base64,{src['data']}"},
            })
        # document-блоки сюда не попадают: PDF в openai-режиме уже сконвертирован в картинки
    return out


def call_claude(api_key: str, base_url: str, auth_mode: str, model: str,
                system_prompt: str, content_blocks: list,
                api_format: str = "anthropic",
                max_tokens: int = 8000, timeout: int = 300) -> str:
    """Прямой HTTP-вызов API. Поддерживает формат Anthropic (/v1/messages)
    и OpenAI-совместимый (/v1/chat/completions) — реселлеры используют оба."""
    url = base_url.rstrip("/")

    headers = {"content-type": "application/json"}
    if auth_mode == "Authorization: Bearer":
        headers["authorization"] = f"Bearer {api_key}"
    else:
        headers["x-api-key"] = api_key

    if api_format == "openai":
        if not url.endswith("/chat/completions"):
            url += "/v1/chat/completions" if not url.endswith("/v1") else "/chat/completions"
        payload = {
            "model": model,
            "max_tokens": max_tokens,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": to_openai_content(content_blocks)},
            ],
        }
    else:
        if not url.endswith("/v1/messages"):
            url += "/v1/messages" if not url.endswith("/v1") else "/messages"
        headers["anthropic-version"] = API_VERSION
        payload = {
            "model": model,
            "max_tokens": max_tokens,
            "system": system_prompt,
            "messages": [{"role": "user", "content": content_blocks}],
        }

    last_err = None
    for attempt in range(3):
        try:
            resp = httpx.post(url, headers=headers, json=payload, timeout=timeout)
            if resp.status_code == 200:
                data = resp.json()
                if api_format == "openai":
                    content = data["choices"][0]["message"]["content"]
                    if isinstance(content, list):  # некоторые прокси возвращают список блоков
                        content = "".join(c.get("text", "") for c in content)
                    return content or ""
                return "".join(
                    b.get("text", "") for b in data.get("content", []) if b.get("type") == "text"
                )
            if resp.status_code in (429, 500, 502, 503, 529):
                last_err = f"HTTP {resp.status_code}: {resp.text[:300]}"
                time.sleep(3 * (attempt + 1))
                continue
            raise RuntimeError(f"Ошибка API (HTTP {resp.status_code}): {resp.text[:500]}")
        except httpx.TimeoutException:
            last_err = "Таймаут запроса"
            time.sleep(2)
        except httpx.HTTPError as e:
            last_err = str(e)
            time.sleep(2)
    raise RuntimeError(f"Не удалось получить ответ от API после 3 попыток. Последняя ошибка: {last_err}")


def parse_json_response(raw: str) -> dict:
    """Достаёт JSON из ответа модели (на случай markdown-ограждений или преамбулы)."""
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end > start:
            return json.loads(text[start:end + 1])
        raise


# =========================== ОТЧЁТ И АНАЛИТИКА ===========================

VERDICT_RU = {"correct": "✅ Верно", "partial": "🟡 Частично", "wrong": "❌ Неверно", "skipped": "⬜ Пропущено"}


def tasks_dataframe(result: dict) -> pd.DataFrame:
    rows = []
    for t in result.get("tasks", []):
        rows.append({
            "№": str(t.get("number", "")),
            "Тема": t.get("topic", ""),
            "Вердикт": VERDICT_RU.get(t.get("verdict"), t.get("verdict", "")),
            "Баллы": t.get("score", 0),
            "Макс.": t.get("max_score", 0),
            "Тип ошибки": t.get("error_type") or "—",
            "Комментарий": t.get("comment", ""),
            "Правильное решение": t.get("correct_solution", ""),
        })
    return pd.DataFrame(rows)


def build_markdown_report(result: dict, subject: str) -> str:
    lines = [
        f"# Отчёт о проверке работы — {subject}",
        "",
        f"**Ученик:** {result.get('student_name') or 'не указан'}  ",
        f"**Работа:** {result.get('work_title') or '—'}  ",
        f"**Оценка:** {result.get('grade_5', '—')} ({result.get('overall_score', '—')}/100)",
        "",
        "## Общий вывод",
        result.get("summary", ""),
        "",
        "## Разбор заданий",
    ]
    for t in result.get("tasks", []):
        lines += [
            f"### Задание {t.get('number', '')} — {t.get('topic', '')}",
            f"- **Вердикт:** {VERDICT_RU.get(t.get('verdict'), t.get('verdict'))}"
            f" ({t.get('score', 0)}/{t.get('max_score', 0)} баллов)",
            f"- **Ответ ученика:** {t.get('student_answer', '')}",
            f"- **Тип ошибки:** {t.get('error_type') or '—'}",
            f"- **Комментарий:** {t.get('comment', '')}",
            f"- **Правильное решение:** {t.get('correct_solution', '')}",
            "",
        ]
    if result.get("strengths"):
        lines += ["## Сильные стороны"] + [f"- {s}" for s in result["strengths"]] + [""]
    if result.get("weaknesses"):
        lines += ["## Слабые места"] + [f"- {s}" for s in result["weaknesses"]] + [""]
    if result.get("recommendations"):
        lines += ["## Рекомендации"] + [f"- {s}" for s in result["recommendations"]] + [""]
    if result.get("teacher_comment"):
        lines += ["## Комментарий преподавателя", result["teacher_comment"], ""]
    return "\n".join(lines)


# ============================== ИНТЕРФЕЙС ==============================

st.set_page_config(page_title=APP_TITLE, page_icon="📝", layout="wide")
st.title("📝 " + APP_TITLE)
st.caption("Загрузите работу (фото, PDF или DOCX) — получите оценку, разбор ошибок, комментарии и аналитику.")

with st.sidebar:
    st.header("⚙️ Подключение к API")
    api_key = st.text_input("API-ключ", type="password",
                            help="Ключ Claude API (в т.ч. от реселлера)")
    base_url = st.text_input("Base URL", value="https://api.anthropic.com",
                             help="Эндпоинт реселлера, напр. https://api.example-reseller.com")
    auth_mode = st.selectbox("Формат авторизации", ["x-api-key", "Authorization: Bearer"],
                             help="Официальный API использует x-api-key; многие реселлеры — Bearer.")
    api_format_label = st.selectbox(
        "Формат API",
        ["Anthropic (/v1/messages)", "OpenAI-совместимый (/v1/chat/completions)"],
        help="Если модель пишет, что «не видит изображения» — реселлер, скорее всего, "
             "проксирует OpenAI-формат и теряет картинки. Переключитесь на OpenAI-совместимый.",
    )
    api_format = "openai" if "OpenAI" in api_format_label else "anthropic"
    model_choice = st.selectbox("Модель", DEFAULT_MODELS + ["Другая…"])
    model = st.text_input("Название модели", value="") if model_choice == "Другая…" else model_choice

    if st.button("👁 Тест подключения и зрения", use_container_width=True,
                 help="Отправляет крошечную картинку и проверяет, видит ли её модель."):
        if not api_key or not model.strip():
            st.error("Сначала укажите ключ и модель.")
        else:
            img = Image.new("RGB", (200, 200), (255, 255, 255))
            for x in range(50, 150):
                for y in range(50, 150):
                    img.putpixel((x, y), (220, 30, 30))
            buf = io.BytesIO()
            img.save(buf, format="JPEG")
            b64 = base64.b64encode(buf.getvalue()).decode()
            test_blocks = [
                {"type": "text", "text": "Ответь одним-двумя словами: какого цвета квадрат на картинке?"},
                {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": b64}},
            ]
            try:
                with st.spinner("Проверяю…"):
                    reply = call_claude(api_key, base_url, auth_mode, model.strip(),
                                        "Ты видишь изображения. Отвечай кратко.",
                                        test_blocks, api_format=api_format,
                                        max_tokens=50, timeout=60)
                if any(w in reply.lower() for w in ("красн", "red")):
                    st.success(f"✅ Зрение работает. Ответ модели: {reply}")
                else:
                    st.error(f"⚠️ Модель ответила, но картинку, похоже, не видит: «{reply}». "
                             "Попробуйте другой формат API или другую модель.")
            except Exception as e:
                st.error(f"Ошибка подключения: {e}")

    st.divider()
    st.header("📚 Параметры проверки")
    subject = st.selectbox("Предмет", list(SUBJECTS.keys()))
    grade_level = st.selectbox("Класс / уровень", [
        "1–4 класс", "5–6 класс", "7–8 класс", "9 класс", "10–11 класс", "Студент / взрослый",
    ], index=3)
    strictness = st.select_slider("Строгость", ["Мягкая", "Стандартная", "Строгая (экзаменационная)"],
                                  value="Стандартная")
    extra_instructions = st.text_area(
        "Доп. указания (необязательно)",
        placeholder="Например: критерии оценивания, правильные ответы, на что обратить внимание…",
    )

uploaded_files = st.file_uploader(
    "Файлы работы ученика",
    type=["jpg", "jpeg", "png", "webp", "gif", "bmp", "pdf", "docx", "doc"],
    accept_multiple_files=True,
)

if uploaded_files:
    imgs = [f for f in uploaded_files if f.name.lower().rsplit(".", 1)[-1]
            in ("jpg", "jpeg", "png", "webp", "gif", "bmp")]
    if imgs:
        cols = st.columns(min(4, len(imgs)))
        for i, f in enumerate(imgs):
            with cols[i % len(cols)]:
                st.image(f, caption=f.name, use_container_width=True)

run = st.button("🚀 Проверить работу", type="primary", use_container_width=True)

if run:
    if not api_key:
        st.error("Введите API-ключ в боковой панели.")
        st.stop()
    if not model.strip():
        st.error("Укажите название модели.")
        st.stop()
    if not uploaded_files:
        st.error("Загрузите хотя бы один файл с работой.")
        st.stop()

    blocks, warnings = build_content_blocks(uploaded_files, api_format=api_format)
    for w in warnings:
        st.warning(w)
    if not blocks:
        st.error("Ни один файл не удалось подготовить к проверке.")
        st.stop()

    grading_note = ("5 = 90-100, 4 = 75-89, 3 = 55-74, 2 = меньше 55"
                    if strictness != "Строгая (экзаменационная)"
                    else "5 = 93-100, 4 = 80-92, 3 = 60-79, 2 = меньше 60")
    system_prompt = SYSTEM_PROMPT.format(
        subject=subject, subject_hint=SUBJECTS[subject],
        grade_level=grade_level, strictness=strictness, grading_note=grading_note,
    )

    task_text = "Проверь приложенную работу ученика и верни результат строго по указанной JSON-схеме."
    if extra_instructions.strip():
        task_text += f"\n\nДополнительные указания преподавателя:\n{extra_instructions.strip()}"
    blocks.append({"type": "text", "text": task_text})

    with st.spinner("Проверяю работу… Это может занять до пары минут."):
        try:
            raw = call_claude(api_key, base_url, auth_mode, model.strip(),
                              system_prompt, blocks, api_format=api_format)
            result = parse_json_response(raw)
        except json.JSONDecodeError:
            st.error("Модель вернула ответ не в формате JSON. Попробуйте ещё раз или смените модель.")
            with st.expander("Сырой ответ модели"):
                st.code(raw)
            st.stop()
        except Exception as e:
            st.error(f"Ошибка: {e}")
            st.stop()

    st.session_state["result"] = result
    st.session_state["subject"] = subject

# ============================ ВЫВОД РЕЗУЛЬТАТА ============================

if "result" in st.session_state:
    result = st.session_state["result"]
    subject = st.session_state.get("subject", "")
    tasks = result.get("tasks", [])

    st.divider()
    st.subheader("Результат проверки")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Оценка (5-балльная)", result.get("grade_5", "—"))
    c2.metric("Балл (из 100)", result.get("overall_score", "—"))
    total = sum(t.get("max_score", 0) for t in tasks)
    got = sum(t.get("score", 0) for t in tasks)
    c3.metric("Баллы по заданиям", f"{got} / {total}" if total else "—")
    correct_n = sum(1 for t in tasks if t.get("verdict") == "correct")
    c4.metric("Верных заданий", f"{correct_n} из {len(tasks)}" if tasks else "—")

    if result.get("student_name") or result.get("work_title"):
        st.caption(f"Ученик: {result.get('student_name') or 'не указан'} · "
                   f"Работа: {result.get('work_title') or '—'}")

    tab1, tab2, tab3, tab4 = st.tabs(["📋 Разбор заданий", "📊 Аналитика", "💬 Комментарии", "📄 Отчёт"])

    with tab1:
        st.write(result.get("summary", ""))
        df = tasks_dataframe(result)
        if not df.empty:
            st.dataframe(df, use_container_width=True, hide_index=True)
            for t in tasks:
                icon = VERDICT_RU.get(t.get("verdict"), "")
                with st.expander(f"{icon} — Задание {t.get('number')} · {t.get('topic', '')} "
                                 f"({t.get('score', 0)}/{t.get('max_score', 0)})"):
                    st.markdown(f"**Ответ ученика:** {t.get('student_answer', '')}")
                    if t.get("error_type"):
                        st.markdown(f"**Тип ошибки:** {t['error_type']}")
                    st.markdown(f"**Комментарий:** {t.get('comment', '')}")
                    st.markdown(f"**Правильное решение:** {t.get('correct_solution', '')}")

    with tab2:
        if tasks:
            left, right = st.columns(2)
            with left:
                st.markdown("**Распределение вердиктов**")
                verdict_counts = pd.Series(
                    [VERDICT_RU.get(t.get("verdict"), "другое") for t in tasks]
                ).value_counts()
                st.bar_chart(verdict_counts)
            with right:
                st.markdown("**Ошибки по типам**")
                err = pd.Series([t.get("error_type") for t in tasks
                                 if t.get("error_type")]).value_counts()
                if not err.empty:
                    st.bar_chart(err)
                else:
                    st.info("Ошибок по типам не зафиксировано 🎉")

            st.markdown("**Баллы по заданиям**")
            score_df = pd.DataFrame({
                "Задание": [str(t.get("number")) for t in tasks],
                "Набрано": [t.get("score", 0) for t in tasks],
                "Максимум": [t.get("max_score", 0) for t in tasks],
            }).set_index("Задание")
            st.bar_chart(score_df)

            if total:
                st.progress(min(1.0, got / total),
                            text=f"Выполнение работы: {round(100 * got / total)}%")
        else:
            st.info("Модель не выделила отдельные задания в работе.")

    with tab3:
        if result.get("strengths"):
            st.markdown("#### 💪 Сильные стороны")
            for s in result["strengths"]:
                st.markdown(f"- {s}")
        if result.get("weaknesses"):
            st.markdown("#### 🎯 Над чем поработать")
            for s in result["weaknesses"]:
                st.markdown(f"- {s}")
        if result.get("recommendations"):
            st.markdown("#### 📌 Рекомендации")
            for s in result["recommendations"]:
                st.markdown(f"- {s}")
        if result.get("teacher_comment"):
            st.markdown("#### 🧑‍🏫 Комментарий преподавателя")
            st.info(result["teacher_comment"])

    with tab4:
        report_md = build_markdown_report(result, subject)
        st.download_button("⬇️ Скачать отчёт (Markdown)", report_md,
                           file_name="report.md", mime="text/markdown",
                           use_container_width=True)
        st.download_button("⬇️ Скачать результат (JSON)",
                           json.dumps(result, ensure_ascii=False, indent=2),
                           file_name="result.json", mime="application/json",
                           use_container_width=True)
        df = tasks_dataframe(result)
        if not df.empty:
            st.download_button("⬇️ Скачать таблицу заданий (CSV)",
                               df.to_csv(index=False).encode("utf-8-sig"),
                               file_name="tasks.csv", mime="text/csv",
                               use_container_width=True)
        with st.expander("Предпросмотр отчёта"):
            st.markdown(report_md)
