"""
checker.py — извлечение контента из файлов и проверка работ через Google Gemini API.

Ключевые особенности:
- Проверка работ через Gemini (google-genai). Vision — основной путь: фото/скан
  отправляется напрямую в модель, которая отлично читает печатный и рукописный
  текст, включая кириллицу. Локальный OCR больше не нужен.
- Надёжный парсинг ответа модели через JSON (+ защитный построчный fallback-парсер).
- Запрос к модели в формате structured JSON.
- Защита от пустого текста из сканов-PDF: при пустом текстовом слое PDF
  рендерится в изображение и отправляется в Gemini как картинка.
- 6 предметов: Английский, Русский, Биология, Математика, Физика, Химия.
- Модель по умолчанию: gemini-2.5-flash (доступна на free-tier), с возможностью
  переключиться на gemini-2.5-pro в интерфейсе.
"""

import json
import re
from io import BytesIO

from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document
from PIL import Image, ImageOps

# --- Конфигурация ---
# Для vision-моделей оптимальна сторона ~1568px по длинной стороне:
# хорошая читаемость рукописного текста без перерасхода токенов.
MAX_IMAGE_SIDE = 1568
JPEG_QUALITY = 90

# Доступные модели Gemini (для выбора в UI). Flash доступна на бесплатном тарифе.
AVAILABLE_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-pro",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
]
DEFAULT_MODEL = "gemini-2.5-flash"
# Лимит на ТЕКСТ ответа. У Gemini 2.5 есть отдельные «thinking»-токены, которые
# раньше съедали этот лимит и обрывали JSON на середине (из-за чего ошибки
# «терялись» и работа казалась правильной). Теперь лимит щедрый, а thinking
# ограничен отдельно (см. THINKING_BUDGET), чтобы не воровать токены у ответа.
MAX_TOKENS = 8192
THINKING_BUDGET = 512  # небольшой бюджет на рассуждения; 0 = полностью выключить
MAX_RETRIES = 3        # повторы при временных сбоях (503/429/пустой ответ)

# Минимальная длина текста, ниже которой PDF/DOCX считаем «пустым» (скан/картинка)
MIN_TEXT_LEN = 15


# ============================================================
#  Подготовка изображения
# ============================================================
def _prep_image(img: Image.Image) -> Image.Image:
    """Корректирует ориентацию по EXIF и масштабирует под vision."""
    try:
        img = ImageOps.exif_transpose(img)
    except Exception:
        pass
    img = img.convert("RGB")
    if max(img.size) > MAX_IMAGE_SIDE:
        ratio = MAX_IMAGE_SIDE / max(img.size)
        img = img.resize(
            (max(1, int(img.width * ratio)), max(1, int(img.height * ratio))),
            Image.LANCZOS,
        )
    return img


def _img_to_bytes(img: Image.Image) -> bytes:
    """Возвращает JPEG-байты изображения (для отправки в Gemini)."""
    buf = BytesIO()
    img.save(buf, format="JPEG", quality=JPEG_QUALITY)
    return buf.getvalue()


# ============================================================
#  Извлечение контента из загруженного файла
# ============================================================
def extract_content(file) -> dict:
    """
    Возвращает словарь одного из видов:
      {"type": "text", "data": <str>}
      {"type": "image", "data": <bytes>, "media_type": <str>}

    Гарантированно сбрасывает указатель файла перед чтением,
    чтобы повторное чтение (превью + обработка) не давало пустоту.
    """
    ext = file.name.rsplit(".", 1)[-1].lower()

    # Всегда читаем сырые байты один раз — это устойчиво к положению указателя.
    try:
        file.seek(0)
    except Exception:
        pass
    raw = file.read()

    # ---------- PDF ----------
    if ext == "pdf":
        text = ""
        try:
            reader = PdfReader(BytesIO(raw))
            text = "\n".join((p.extract_text() or "") for p in reader.pages).strip()
        except Exception:
            text = ""

        if len(text) >= MIN_TEXT_LEN:
            return {"type": "text", "data": text}

        # PDF без текстового слоя (скан) — рендерим первую страницу в картинку
        img = _render_pdf_first_page(raw)
        if img is not None:
            img = _prep_image(img)
            return {
                "type": "image",
                "data": _img_to_bytes(img),
                "media_type": "image/jpeg",
            }
        # Если рендер недоступен — возвращаем то, что есть
        return {"type": "text",
                "data": text or "[PDF не содержит распознаваемого текста]"}

    # ---------- DOCX ----------
    if ext == "docx":
        try:
            doc = Document(BytesIO(raw))
            parts = [p.text for p in doc.paragraphs]
            # Текст из таблиц тоже учитываем
            for table in doc.tables:
                for r in table.rows:
                    parts.extend(c.text for c in r.cells)
            text = "\n".join(t for t in parts if t).strip()
        except Exception:
            text = ""
        return {"type": "text", "data": text or "[DOCX пустой или не читается]"}

    # ---------- Изображения ----------
    img = Image.open(BytesIO(raw))
    img = _prep_image(img)
    return {
        "type": "image",
        "data": _img_to_bytes(img),
        "media_type": "image/jpeg",  # сохраняем как JPEG после обработки
    }


def _render_pdf_first_page(raw: bytes):
    """Рендерит первую страницу PDF в PIL.Image (PyMuPDF или pdf2image)."""
    # 1) PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=raw, filetype="pdf")
        page = doc.load_page(0)
        pix = page.get_pixmap(dpi=200)
        return Image.open(BytesIO(pix.tobytes("png")))
    except Exception:
        pass
    # 2) pdf2image (требует poppler)
    try:
        from pdf2image import convert_from_bytes
        pages = convert_from_bytes(raw, dpi=200, first_page=1, last_page=1)
        if pages:
            return pages[0]
    except Exception:
        pass
    return None


# ============================================================
#  Конфигурация предметов (поля и критерии)
# ============================================================
SUBJECT_CONFIG = {
    "Английский язык": {
        "role": "опытный учитель английского языка",
        "scores": [
            ("GRAMMAR_SCORE", "Грамматика"),
            ("VOCABULARY_SCORE", "Лексика"),
            ("SPELLING_SCORE", "Орфография"),
            ("COHERENCE_SCORE", "Связность"),
        ],
        "errors": [
            ("GRAMMAR_ERRORS", "🔴 Грамматические ошибки"),
            ("SPELLING_ERRORS", "🟠 Орфографические ошибки"),
            ("VOCABULARY_ERRORS", "🟡 Лексические ошибки"),
        ],
        "level_hint": "уровень владения языком: A1/A2/B1/B2/C1/C2",
    },
    "Русский язык": {
        "role": "опытный учитель русского языка и литературы",
        "scores": [
            ("SPELLING_SCORE", "Орфография"),
            ("PUNCTUATION_SCORE", "Пунктуация"),
            ("GRAMMAR_SCORE", "Грамматика"),
            ("STYLE_SCORE", "Стиль и связность"),
        ],
        "errors": [
            ("SPELLING_ERRORS", "🔴 Орфографические ошибки"),
            ("PUNCTUATION_ERRORS", "🟠 Пунктуационные ошибки"),
            ("STYLE_ERRORS", "🟡 Стилистические/речевые ошибки"),
        ],
        "level_hint": "уровень грамотности: Начальный/Базовый/Уверенный/Высокий",
    },
    "Биология": {
        "role": "опытный учитель биологии",
        "scores": [
            ("KNOWLEDGE_SCORE", "Знание"),
            ("CORRECTNESS_SCORE", "Термины"),
            ("LOGIC_SCORE", "Логика"),
            ("EXAMPLES_SCORE", "Примеры"),
        ],
        "errors": [
            ("FACTUAL_ERRORS", "🔴 Фактические ошибки"),
            ("TERM_ERRORS", "🟠 Ошибки в терминах"),
            ("STRUCTURE_ERRORS", "🟡 Ошибки в логике"),
        ],
        "level_hint": "уровень знаний: Начальный/Базовый/Продвинутый/Эксперт",
    },
    "Математика": {
        "role": "опытный учитель математики",
        "scores": [
            ("CALCULATION_SCORE", "Вычисления"),
            ("LOGIC_SCORE", "Логика"),
            ("METHOD_SCORE", "Метод"),
            ("PRESENTATION_SCORE", "Оформление"),
        ],
        "errors": [
            ("CALCULATION_ERRORS", "🔴 Вычислительные ошибки"),
            ("LOGIC_ERRORS", "🟠 Логические ошибки"),
            ("METHOD_ERRORS", "🟡 Ошибки в методе"),
        ],
        "level_hint": "уровень: Начальный/Базовый/Продвинутый/Эксперт",
    },
    "Физика": {
        "role": "опытный учитель физики",
        "scores": [
            ("CONCEPT_SCORE", "Понимание законов"),
            ("CALCULATION_SCORE", "Вычисления"),
            ("UNITS_SCORE", "Единицы измерения"),
            ("PRESENTATION_SCORE", "Оформление"),
        ],
        "errors": [
            ("CONCEPT_ERRORS", "🔴 Ошибки в законах/понятиях"),
            ("CALCULATION_ERRORS", "🟠 Вычислительные ошибки"),
            ("UNITS_ERRORS", "🟡 Ошибки в единицах измерения"),
        ],
        "level_hint": "уровень: Начальный/Базовый/Продвинутый/Эксперт",
    },
    "Химия": {
        "role": "опытный учитель химии",
        "scores": [
            ("EQUATION_SCORE", "Уравнения реакций"),
            ("CALCULATION_SCORE", "Расчёты"),
            ("TERM_SCORE", "Термины и понятия"),
            ("PRESENTATION_SCORE", "Оформление"),
        ],
        "errors": [
            ("EQUATION_ERRORS", "🔴 Ошибки в уравнениях реакций"),
            ("CALCULATION_ERRORS", "🟠 Расчётные ошибки"),
            ("TERM_ERRORS", "🟡 Ошибки в терминах/формулах"),
        ],
        "level_hint": "уровень: Начальный/Базовый/Продвинутый/Эксперт",
    },
}


def get_subjects() -> list:
    return list(SUBJECT_CONFIG.keys())


def get_subject_config(subject: str) -> dict:
    return SUBJECT_CONFIG.get(subject, SUBJECT_CONFIG["Английский язык"])


def get_models() -> list:
    return list(AVAILABLE_MODELS)


# ============================================================
#  Построение промпта (просим строго JSON)
# ============================================================
def build_prompt(subject: str) -> str:
    cfg = get_subject_config(subject)
    score_keys = [k for k, _ in cfg["scores"]]
    error_keys = [k for k, _ in cfg["errors"]]

    score_lines = "\n".join(f'  "{k}": <целое число от 1 до 10>,' for k in score_keys)
    error_lines = "\n".join(
        f'  "{k}": "<краткий список ошибок через точку с запятой, или \\"нет\\">",'
        for k in error_keys
    )

    return f"""Ты — {cfg['role']} и строгий, внимательный проверяющий. Проверь домашнюю работу ученика и дай детальный, объективный анализ на русском языке.

ПОРЯДОК РАБОТЫ:
1. Если работа на изображении — сначала дословно распознай весь текст (включая рукописный), СОХРАНЯЯ все ошибки ученика КАК ЕСТЬ (не исправляй их при чтении!).
2. Затем внимательно, предложение за предложением, найди ВСЕ ошибки без исключения.

КРИТИЧЕСКИ ВАЖНО:
- Будь придирчив: ученики ПОЧТИ ВСЕГДА делают ошибки. Если ты не нашёл ни одной ошибки — скорее всего ты проверил невнимательно. Перепроверь ещё раз.
- Пиши "нет" ТОЛЬКО если ты ДЕЙСТВИТЕЛЬНО уверен, что ошибок данного типа нет.
- Для каждой ошибки укажи конкретно: что написано неправильно и как должно быть (напр. "goes → go").
- Оценка должна ОТРАЖАТЬ количество ошибок: много ошибок → низкая оценка.

Ответь СТРОГО валидным JSON-объектом без markdown, без пояснений до или после. Структура:
{{
  "GRADE": <итоговая оценка, целое число от 1 до 10>,
{score_lines}
{error_lines}
  "STRENGTHS": "<что сделано хорошо, через точку с запятой>",
  "RECOMMENDATIONS": "<конкретные рекомендации по улучшению, через точку с запятой>",
  "LEVEL": "<{cfg['level_hint']}>",
  "SUMMARY": "<краткое общее заключение в 2-3 предложениях>"
}}

Важно: значения ошибок — это строки. Если ошибок нет — поставь "нет". Не добавляй никаких полей, кроме перечисленных. Верни ТОЛЬКО JSON."""


# ============================================================
#  Парсинг ответа модели
# ============================================================
def _extract_json_block(text: str):
    """Извлекает первый JSON-объект из текста (на случай обёрток/markdown)."""
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if fenced:
        try:
            return json.loads(fenced.group(1))
        except Exception:
            pass
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except Exception:
            pass
    return None


def _fallback_line_parse(text: str, subject: str) -> dict:
    """Запасной построчный парсер «КЛЮЧ: значение» (многострочно-устойчивый)."""
    cfg = get_subject_config(subject)
    known = {"GRADE", "STRENGTHS", "RECOMMENDATIONS", "LEVEL", "SUMMARY"}
    known |= {k for k, _ in cfg["scores"]}
    known |= {k for k, _ in cfg["errors"]}

    def clean(s: str) -> str:
        s = s.strip()
        s = re.sub(r"[*_`#>]+", "", s)
        return s.strip().strip('"').strip()

    result = {}
    current_key = None
    for raw_line in text.splitlines():
        line = raw_line.strip().lstrip("*#-• ").strip()
        m = re.match(r"^[*_`]*([A-Z_]+)[*_`]*\s*[:：]\s*(.*)$", line)
        if m and m.group(1) in known:
            current_key = m.group(1)
            result[current_key] = clean(m.group(2))
        elif current_key and line:
            result[current_key] = (result[current_key] + " " + clean(line)).strip()
    return result


def _coerce_int(val, default=0) -> int:
    try:
        m = re.search(r"-?\d+", str(val))
        return int(m.group()) if m else default
    except Exception:
        return default


def parse_response(text: str, subject: str) -> dict:
    """Возвращает нормализованный словарь результата проверки."""
    cfg = get_subject_config(subject)
    data = _extract_json_block(text)
    if data is None or not isinstance(data, dict):
        data = _fallback_line_parse(text, subject)

    if not data:
        return {"_raw": text, "_parse_failed": True}

    result = {}
    result["GRADE"] = _coerce_int(data.get("GRADE"), 0)
    for k, _ in cfg["scores"]:
        result[k] = _coerce_int(data.get(k), 0)
    for k, _ in cfg["errors"]:
        result[k] = str(data.get(k, "нет")).strip() or "нет"
    for k in ("STRENGTHS", "RECOMMENDATIONS", "LEVEL", "SUMMARY"):
        result[k] = str(data.get(k, "—")).strip() or "—"
    return result


# ============================================================
#  Работа с Gemini
# ============================================================
def _make_client(api_key: str):
    return genai.Client(api_key=api_key)


def _build_config(json_mode: bool):
    """Конфиг генерации с бюджетом thinking и (опционально) JSON-режимом."""
    kwargs = dict(max_output_tokens=MAX_TOKENS, temperature=0.2)
    if json_mode:
        kwargs["response_mime_type"] = "application/json"
    # Ограничиваем thinking отдельным бюджетом, чтобы он не съедал токены
    # ответа и не обрывал JSON. Поддерживается не всеми версиями SDK/моделей.
    try:
        kwargs["thinking_config"] = types.ThinkingConfig(
            thinking_budget=THINKING_BUDGET)
    except Exception:
        pass
    try:
        return types.GenerateContentConfig(**kwargs)
    except Exception:
        # Если thinking_config не поддерживается — убираем его и пробуем снова
        kwargs.pop("thinking_config", None)
        return types.GenerateContentConfig(**kwargs)


class GeminiError(Exception):
    """Ошибка запроса к Gemini (понятное сообщение для пользователя)."""


def _ask_gemini(client, model: str, parts: list, json_mode: bool = True) -> str:
    """
    Отправляет запрос в Gemini с повторами при временных сбоях.
    Бросает GeminiError с понятным текстом, если все попытки провалились
    — чтобы приложение НИКОГДА не показывало «всё правильно» вместо ошибки.
    """
    import time
    last_err = None
    for attempt in range(MAX_RETRIES):
        try:
            resp = client.models.generate_content(
                model=model, contents=parts, config=_build_config(json_mode),
            )
            text = (getattr(resp, "text", "") or "").strip()

            # Проверяем причину завершения: обрыв по лимиту токенов = неполный JSON
            fr = ""
            try:
                fr = str(resp.candidates[0].finish_reason or "")
            except Exception:
                pass
            if text:
                return text
            # Пустой ответ (напр. весь лимит ушёл в thinking) — повторяем
            last_err = f"пустой ответ модели (finish_reason={fr})"
        except Exception as e:
            msg = str(e)
            last_err = msg
            # 503/429/500 — временные, имеет смысл повторить
            transient = any(c in msg for c in ("503", "429", "500", "UNAVAILABLE",
                                               "RESOURCE_EXHAUSTED", "overloaded"))
            if not transient:
                break  # неповторяемая ошибка (напр. неверный ключ) — не тянем
        # экспоненциальная пауза перед следующей попыткой
        if attempt < MAX_RETRIES - 1:
            time.sleep(1.5 * (attempt + 1))

    raise GeminiError(_friendly_error(last_err or "неизвестная ошибка"))


def _friendly_error(msg: str) -> str:
    """Превращает техническую ошибку в понятное сообщение."""
    low = msg.lower()
    if "429" in msg or "resource_exhausted" in low:
        return ("Исчерпана квота этой модели (429). Переключитесь на другую модель "
                "(напр. gemini-2.5-flash) или подождите и повторите.")
    if "503" in msg or "unavailable" in low or "overloaded" in low:
        return ("Модель временно перегружена (503). Попробуйте повторить через несколько секунд.")
    if "api key" in low or "401" in msg or "403" in msg or "permission" in low:
        return "Проблема с API-ключом. Проверьте ключ Gemini в боковой панели."
    return f"Ошибка запроса к Gemini: {msg[:200]}"


def check_homework(content: dict, api_key: str, subject: str = "Английский язык",
                   model: str = DEFAULT_MODEL, debug: bool = False) -> dict:
    prompt = build_prompt(subject)
    client = _make_client(api_key)
    dbg = {"модель": model}

    if content["type"] == "text":
        body = prompt + "\n\nРабота ученика:\n" + content["data"]
        raw_text = _ask_gemini(client, model, [body])
        dbg["режим"] = "текст"
    else:
        # ОСНОВНОЙ ПУТЬ: отправляем изображение напрямую в Gemini (vision).
        # Gemini хорошо читает печатный и рукописный текст, включая кириллицу.
        img_part = types.Part.from_bytes(
            data=content["data"],
            mime_type=content.get("media_type", "image/jpeg"),
        )
        raw_text = _ask_gemini(client, model, [img_part, prompt])
        dbg["режим"] = "изображение (vision)"
        dbg["размер_изображения_байт"] = len(content.get("data", b""))

    dbg["длина_ответа"] = len(raw_text)
    dbg["ответ_начало"] = raw_text[:400]

    result = parse_response(raw_text, subject)
    result["Предмет"] = subject
    if debug:
        result["_debug"] = dbg
    return result


def diagnose_vision(api_key: str, model: str = DEFAULT_MODEL) -> dict:
    """
    Проверяет, читает ли модель текст с изображения.
    Генерирует картинку с чётким кодом и просит модель его назвать.
    """
    from PIL import ImageDraw, ImageFont
    secret = "TURTLE7"
    img = Image.new("RGB", (400, 160), "white")
    d = ImageDraw.Draw(img)
    try:
        f = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 60)
    except Exception:
        f = ImageFont.load_default()
    d.text((40, 50), secret, fill="black", font=f)
    data = _img_to_bytes(img)

    out = {"secret": secret, "model": model}
    try:
        client = _make_client(api_key)
        part = types.Part.from_bytes(data=data, mime_type="image/jpeg")
        answer = _ask_gemini(
            client, model,
            [part, "Какое слово/код написано на картинке? Ответь только этим словом."],
            json_mode=False,
        )
        out["answer"] = answer
        out["vision_ok"] = secret.lower() in answer.lower()
    except Exception as e:
        out["error"] = f"{type(e).__name__}: {e}"
        out["vision_ok"] = False
    return out


def check_api_key(api_key: str, model: str = DEFAULT_MODEL) -> dict:
    """Быстрая проверка ключа: простой текстовый запрос."""
    out = {"model": model}
    try:
        client = _make_client(api_key)
        ans = _ask_gemini(client, model, ["Ответь одним словом: готово"], json_mode=False)
        out["ok"] = bool(ans)
        out["answer"] = ans[:100]
    except Exception as e:
        out["ok"] = False
        out["error"] = f"{type(e).__name__}: {e}"
    return out
